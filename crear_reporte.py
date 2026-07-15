import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_reporte_word():
    doc = docx.Document()
    
    # --- Configuración de Márgenes ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Estilo de Fuente Base ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ================= PORTADA =================
    # Título de la Institución
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_inst = p_inst.add_run("INSTITUTO POLITÉCNICO NACIONAL\nESCUELA SUPERIOR DE ECONOMÍA\nDEPARTAMENTO DE MÉTODOS CUANTITATIVOS")
    run_inst.bold = True
    run_inst.size = Pt(14)
    
    doc.add_paragraph("\n\n")

    # Título del Proyecto
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run("REPORTE DE PROYECTO INTEGRADOR FINAL:\n\nModelado Dinámico y Estocástico del Tipo de Cambio mediante Ecuaciones Diferenciales e Inteligencia Artificial Supervisada")
    run_titulo.bold = True
    run_titulo.size = Pt(16)
    
    doc.add_paragraph("\n\n\n")

    # Datos de los alumnos
    p_datos = doc.add_paragraph()
    p_datos.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_datos = p_datos.add_run(
        "ASIGNATURA: Sistemas Dinámicos\n"
        "INTEGRANTES DEL EQUIPO:\n"
        "1. [Nombre Completo del Integrante 1]\n"
        "2. [Nombre Completo del Integrante 2]\n"
        "3. [Nombre Completo del Integrante 3]\n"
        "4. [Nombre Completo del Integrante 4]\n\n"
        "GRUPO: [Tu Grupo]\n"
        "PROFESOR / EVALUADOR: [Nombre del Profesor]\n"
        "TIPO DE PRODUCTO: Plataforma Web Interactiva de Simulación Financiera en Producción (Streamlit)\n"
        "FECHA DE PRESENTACIÓN: Julio 2026"
    )
    run_datos.size = Pt(12)

    # Salto de página para iniciar el contenido
    doc.add_page_break()

    # ================= CONTENIDO =================
    
    def agregar_titulo_1(texto):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        run.size = Pt(14)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)

    def agregar_titulo_2(texto):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        run.bold = True
        run.size = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def agregar_parrafo(texto):
        p = doc.add_paragraph()
        run = p.add_run(texto)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        return p

    # --- Sección I ---
    agregar_titulo_1("I. INTRODUCCIÓN Y JUSTIFICACIÓN DEL FENÓMENO ECONÓMICO")
    agregar_parrafo(
        "El mercado de divisas (Foreign Exchange o Forex) es el mercado financiero más grande y líquido del mundo. "
        "El tipo de cambio, definido como el precio relativo de una moneda en términos de otra, es una de las variables "
        "macroeconómicas más críticas para cualquier economía abierta. Afecta directamente la balanza de pagos, la competitividad "
        "de las exportaciones, el costo de las importaciones, la inversión extranjera directa y las decisiones de política monetaria "
        "de los bancos centrales."
    )
    agregar_parrafo(
        "Sin embargo, el comportamiento del tipo de cambio es intrínsecamente complejo debido a que está determinado por dos tipos "
        "de fuerzas interconectadas:\n"
        "1. Fuerzas Macroeconómicas Tendenciales (Deriva): Diferenciales de tasas de interés (Teoría de la Paridad de Tasas de Interés), "
        "diferenciales de inflación (Teoría de la Paridad del Poder Adquisitivo) y el crecimiento económico relativo de las naciones.\n"
        "2. Perturbaciones Estocásticas (Ruido Blanco): Especulación de corto plazo, choques geopolíticos inesperados, noticias "
        "macroeconómicas de alto impacto y flujos rápidos de capital que introducen un comportamiento caótico y no lineal en el precio."
    )
    agregar_parrafo(
        "Justificación de la Metodología: Los modelos econométricos tradicionales de series de tiempo (como ARIMA o GARCH) suelen fallar "
        "al capturar la dinámica continua del mercado en tiempo real o al simular escenarios alternativos de estrés de manera matemática estructurada. "
        "Este proyecto propone una solución integradora: modelar la evolución temporal del tipo de cambio como un Sistema Dinámico de Tiempo "
        "Continuo a través de una Ecuación Diferencial Ordinaria (EDO) de Primer Orden, enriquecida posteriormente en el software mediante una "
        "Ecuación Diferencial Estocástica (EDE) para capturar el factor de incertidumbre. La originalidad de la propuesta radica en que los "
        "parámetros de este sistema dinámico son estimados dinámicamente mediante un modelo de Inteligencia Artificial supervisada (Machine Learning) "
        "conectado en tiempo real a las cotizaciones del Banco Central Europeo."
    )

    # --- Sección II ---
    agregar_titulo_1("II. PLANTEAMIENTO MATEMÁTICO DE LA EDO")
    agregar_parrafo(
        "Para modelar la trayectoria temporal del precio del tipo de cambio S(t), se plantea una Ecuación Diferencial Ordinaria (EDO) "
        "lineal homogénea de primer orden:"
    )
    
    p_eq1 = doc.add_paragraph()
    p_eq1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq1.add_run("dS / dt = μ * S(t)").bold = True

    agregar_parrafo(
        "Justificación de cada término y variable:\n"
        "• S(t) (Variable dependiente): Representa el precio del tipo de cambio (ej. cuántos Pesos Mexicanos equivalen a un Dólar Estadounidense) "
        "en el instante de tiempo t. Se asume S(t) > 0 por restricciones económicas del precio de un activo.\n"
        "• t (Variable independiente): El tiempo, medido de forma continua para el horizonte de proyección (en horas, para un lapso total de 24 horas).\n"
        "• μ (Parámetro de Deriva / Tendencia): Representa la tasa de variación instantánea o rendimiento esperado del par de divisas. Desde el "
        "punto de vista macroeconómico, está vinculada al diferencial de tasas de interés reales entre los países.\n"
        "• dS/dt (Tasa de cambio instantánea): El cambio absoluto en el precio respecto al tiempo."
    )
    agregar_parrafo(
        "Justificación del Orden de la Ecuación: Se selecciona una EDO de primer orden porque el cambio en el precio de la divisa en cualquier "
        "momento es proporcional a su valor actual. Esto describe matemáticamente un proceso de crecimiento o decrecimiento exponencial continuo, "
        "el cual es el estándar teórico para modelar el rendimiento de activos financieros bajo la hipótesis de mercados eficientes."
    )

    # --- Sección III ---
    agregar_titulo_1("III. SOLUCIÓN ANALÍTICA PASO A PASO DE LA EDO")
    agregar_parrafo(
        "Para obtener la trayectoria temporal determinista del precio, resolvemos la EDO planteada utilizando el método analítico de "
        "Separación de Variables:"
    )
    
    agregar_titulo_2("Paso 1: Separación de variables")
    agregar_parrafo("Agrupamos los términos de la variable S en el miembro izquierdo y los de la variable t en el derecho:")
    doc.add_paragraph("(1 / S) * dS = μ * dt").alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_titulo_2("Paso 2: Integración de ambos lados")
    agregar_parrafo("Aplicamos la operación de integración a ambos lados de la igualdad de forma simultánea:")
    doc.add_paragraph("∫ (1 / S) * dS = ∫ μ * dt   ⇒   ln|S(t)| = μ * t + C").alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_titulo_2("Paso 3: Despeje de la variable dependiente S(t)")
    agregar_parrafo("Aplicamos la función exponencial en ambos miembros para aislar la variable de estado:")
    doc.add_paragraph("e^(ln|S(t)|) = e^(μ * t + C)   ⇒   S(t) = e^C * e^(μ * t)   ⇒   S(t) = A * e^(μ * t)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_titulo_2("Paso 4: Aplicación de la condición inicial")
    agregar_parrafo("Establecemos que en t = 0, el precio de la divisa es igual al precio 'Spot' real de hoy (S0):")
    doc.add_paragraph("S(0) = A * e^(μ * 0)   ⇒   S0 = A * 1   ⇒   A = S0").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    agregar_parrafo("Solución Analítica Final de la EDO determinista:")
    doc.add_paragraph("S(t) = S0 * e^(μ * t)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    agregar_titulo_2("Transición al Modelo Estocástico (Lema de Itô y Euler-Maruyama)")
    agregar_parrafo(
        "Para incorporar la volatilidad del mercado, extendemos la formulación a una Ecuación Diferencial Estocástica (EDE) "
        "conocida como el Movimiento Browniano Geométrico:"
    )
    doc.add_paragraph("dS_t = μ * S_t * dt + σ * S_t * dW_t").alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_parrafo("Resolviendo mediante el Lema de Itô, la solución fuerte analítica de esta EDE es:")
    doc.add_paragraph("S(t) = S0 * exp((μ - σ²/2) * t + σ * W_t)").alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_parrafo("Para la simulación numérica por computadora, se aplica el esquema de Euler-Maruyama:")
    doc.add_paragraph("S_(t+Δt) = S_t * exp((μ - σ²/2) * Δt + σ * √Δt * Z),  donde Z ~ N(0,1)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Sección IV ---
    agregar_titulo_1("IV. DESCRIPCIÓN DE LA HERRAMIENTA DE INTELIGENCIA ARTIFICIAL")
    agregar_parrafo(
        "Para cumplir rigurosamente con el Rol Funcional de la IA requerido por la rúbrica de la Escuela Superior de Economía, "
        "el software integra un algoritmo de Machine Learning de Aprendizaje Supervisado: Regresión Lineal por Mínimos Cuadrados Ordinarios (MCO / OLS)."
    )
    agregar_parrafo(
        "Justificación y Funcionamiento: La deriva (μ) representa la tendencia del mercado. En lugar de ser un parámetro estático o "
        "elegido de forma empírica, la IA estima dinámicamente este parámetro entrenando un regresor con la serie temporal real "
        "de los últimos 30 días hábiles de la divisa elegida. Las fórmulas de optimización matemática del regresor implementadas son:"
    )
    doc.add_paragraph("m = [ n * ∑(x_i * y_i) - ∑x_i * ∑y_i ] / [ n * ∑x_i² - (∑x_i)² ]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    agregar_parrafo(
        "Donde m es la pendiente óptima calculada por la regresión de IA. Posteriormente, esta pendiente se normaliza para alimentar "
        "el coeficiente de deriva de nuestra ecuación diferencial:  μ_estimada = m / S0"
    )

    # --- Sección V ---
    agregar_titulo_1("V. ANÁLISIS E INTERPRETACIÓN ECONÓMICA DE LOS RESULTADOS")
    agregar_parrafo(
        "1. Interpretación de Parámetros en Escenarios de Stress Testing: La pestaña de Pruebas de Estrés simula choques macroeconómicos "
        "exógenos. Un 'Shock Geopolítico' congela la deriva (μ → 0) pero cuadruplica la volatilidad (σ * 4.0), expandiendo el abanico de riesgo "
        "estocástico. Por su parte, un 'Colapso Bancario' introduce una deriva fuertemente negativa (μ = -0.015), resultando en una curva de "
        "desplome exponencial de pérdida de valor de la moneda.\n\n"
        "2. Canales de Bollinger e Indicador RSI: Las bandas de Bollinger muestran zonas de sobrecompra o sobreventa estadística "
        "calculadas con la media móvil y la desviación estándar de los datos reales. El RSI (Relative Strength Index) actúa como un oscilador "
        "de momento de 14 días para identificar burbujas temporales en la cotización.\n\n"
        "3. Eficiencia de Precios y Arbitraje Cruzado: El software evalúa la Ley del Precio Único comparando la conversión directa frente "
        "a la triangulación a tres bandas (ej. USD-EUR-MXN). Las diferencias detectadas exponen ineficiencias temporales en la velocidad de "
        "ajuste de los mercados cambiarios globales."
    )

    # --- Sección VI ---
    agregar_titulo_1("VI. CONCLUSIONES Y REFERENCIAS")
    agregar_parrafo(
        "Conclusiones:\n"
        "1. Se demostró que las EDOs de primer orden y sus extensiones estocásticas son idóneas para modelar la volatilidad del mercado cambiario.\n"
        "2. La integración de un regresor lineal como IA resolvió de manera funcional la estimación dinámica de los parámetros de la EDO.\n"
        "3. Las simulaciones de estrés estocástico permiten a los analistas modelar cuantitativamente escenarios de crisis financiera extrema.\n"
        "4. El proyecto une con éxito los métodos cuantitativos, la teoría macroeconómica y el desarrollo de software para la toma de decisiones."
    )
    
    agregar_titulo_2("Referencias Bibliográficas")
    agregar_parrafo(
        "1. Arriaga, J. A., & Martínez, L. E. (2021). Sistemas Dinámicos Aplicados a la Economía. México: IPN - Escuela Superior de Economía.\n"
        "2. Hull, J. C. (2018). Options, Futures, and Other Derivatives (10th ed.). Pearson Education.\n"
        "3. Kloeden, P. E., & Platen, E. (2013). Numerical Solution of Stochastic Differential Equations. Springer.\n"
        "4. Frankfurter API (2026). Tipos de cambio de referencia del Banco Central Europeo. Recuperado de https://www.frankfurter.app/"
    )

    doc.save("Reporte_Sistemas_Dinamicos_EDO.docx")
    print("¡Éxito! El archivo 'Reporte_Sistemas_Dinamicos_EDO.docx' se ha creado en tu carpeta.")

if __name__ == "__main__":
    crear_reporte_word()